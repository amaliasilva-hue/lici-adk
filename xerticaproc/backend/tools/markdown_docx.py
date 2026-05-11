"""Conversor Markdown → DOCX usando python-docx (sem depender de pandoc).

Suporta o subconjunto de Markdown produzido pelo redator LLM e pelos
templates determinísticos do Copilot:

- Cabeçalhos (`#` … `######`)
- Parágrafos
- Listas com marcador (`- `, `* `, `+ `) e numeradas (`1. `)
- Negrito `**texto**` / itálico `*texto*` / código inline `` `t` ``
- Tabelas pipe (`| a | b |` + linha separadora `|---|---|`)
- Linhas horizontais (`---`, `***`)
- Blocos de citação (`> `)
- Links `[texto](url)` (renderizados como texto + URL entre parênteses)

Não suporta blocos de código com cercas (raros nos documentos do Copilot)
e imagens — esses casos caem para texto puro preservado.
"""
from __future__ import annotations

import io
import re
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_CODE_RE = re.compile(r"`([^`]+)`")
_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def _add_inline_runs(paragraph, text: str) -> None:
    """Adiciona runs preservando bold/italic/code/links de forma simples."""
    # Estratégia: tokenizar por padrões em ordem de prioridade.
    # Para manter robustez, fazemos passes sequenciais.
    # Substituímos links primeiro para evitar conflito com colchetes.
    text = _LINK_RE.sub(lambda m: f"{m.group(1)} ({m.group(2)})", text)

    # Tokenize sequência preservando marcadores
    pos = 0
    pattern = re.compile(
        r"(\*\*(?P<b>.+?)\*\*)|(`(?P<c>[^`]+)`)|(\*(?P<i>[^*]+)\*)"
    )
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group("b") is not None:
            run = paragraph.add_run(m.group("b"))
            run.bold = True
        elif m.group("c") is not None:
            run = paragraph.add_run(m.group("c"))
            run.font.name = "Consolas"
        elif m.group("i") is not None:
            run = paragraph.add_run(m.group("i"))
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _is_table_separator(line: str) -> bool:
    s = line.strip()
    if not (s.startswith("|") and s.endswith("|")):
        return False
    cells = [c.strip() for c in s.strip("|").split("|")]
    if not cells:
        return False
    return all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c)


def _split_pipe_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _emit_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    cols = max(len(header), max((len(r) for r in rows), default=0))
    if cols == 0:
        return
    tbl = doc.add_table(rows=1 + len(rows), cols=cols)
    tbl.style = "Light Grid Accent 1"
    # header
    for j in range(cols):
        cell = tbl.rows[0].cells[j]
        cell.text = ""
        text = header[j] if j < len(header) else ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
    # body
    for i, row in enumerate(rows, start=1):
        for j in range(cols):
            cell = tbl.rows[i].cells[j]
            cell.text = ""
            text = row[j] if j < len(row) else ""
            _add_inline_runs(cell.paragraphs[0], text)


def _emit_paragraph(doc: Document, text: str, *, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    _add_inline_runs(p, text)


def markdown_to_docx_bytes(md: str, *, title: str | None = None) -> bytes:
    """Converte Markdown em DOCX e devolve os bytes do arquivo .docx."""
    doc = Document()

    # Ajuste de fonte padrão para algo legível tipo "documento oficial".
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if title:
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    lines = md.splitlines()
    i = 0
    n = len(lines)

    def _is_bullet(line: str) -> tuple[bool, str]:
        m = re.match(r"^\s*[-*+]\s+(.*)$", line)
        if m:
            return True, m.group(1)
        return False, ""

    def _is_numbered(line: str) -> tuple[bool, str]:
        m = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m:
            return True, m.group(1)
        return False, ""

    while i < n:
        raw = lines[i]
        line = raw.rstrip()

        # blank line
        if not line.strip():
            i += 1
            continue

        # horizontal rule
        if re.fullmatch(r"\s*([-*_])\1{2,}\s*", line):
            doc.add_paragraph().add_run().add_break()
            i += 1
            continue

        # tabela: linha começa com "|" e a próxima é separator
        if line.lstrip().startswith("|") and i + 1 < n and _is_table_separator(lines[i + 1]):
            header = _split_pipe_row(line)
            i += 2
            rows: list[list[str]] = []
            while i < n and lines[i].lstrip().startswith("|"):
                rows.append(_split_pipe_row(lines[i]))
                i += 1
            _emit_table(doc, header, rows)
            continue

        # cabeçalhos
        m_h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m_h:
            level = len(m_h.group(1))
            text = m_h.group(2).strip()
            # docx suporta Heading 1..9; nível 1 = título principal já posto se houver
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        # citação
        if line.lstrip().startswith("> "):
            buf: list[str] = []
            while i < n and lines[i].lstrip().startswith("> "):
                buf.append(lines[i].lstrip()[2:])
                i += 1
            _emit_paragraph(doc, " ".join(buf), style="Intense Quote")
            continue

        # listas com marcador
        is_b, b_text = _is_bullet(line)
        if is_b:
            while i < n:
                ok, txt = _is_bullet(lines[i])
                if not ok:
                    break
                p = doc.add_paragraph(style="List Bullet")
                _add_inline_runs(p, txt)
                i += 1
            continue

        # listas numeradas
        is_n, n_text = _is_numbered(line)
        if is_n:
            while i < n:
                ok, txt = _is_numbered(lines[i])
                if not ok:
                    break
                p = doc.add_paragraph(style="List Number")
                _add_inline_runs(p, txt)
                i += 1
            continue

        # parágrafo: junta linhas até linha em branco / cabeçalho / lista / tabela
        buf = [line]
        i += 1
        while i < n:
            nxt = lines[i].rstrip()
            if not nxt.strip():
                break
            if re.match(r"^#{1,6}\s+", nxt):
                break
            if _is_bullet(nxt)[0] or _is_numbered(nxt)[0]:
                break
            if nxt.lstrip().startswith("|"):
                break
            if nxt.lstrip().startswith("> "):
                break
            buf.append(nxt)
            i += 1
        _emit_paragraph(doc, " ".join(b.strip() for b in buf))

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
