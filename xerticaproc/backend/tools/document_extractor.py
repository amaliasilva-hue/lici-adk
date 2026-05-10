"""Document extractor — leitura multimodal de anexos (PDF, imagem, DOCX, XLSX).

Estratégia:
  - PDF: tenta PyMuPDF (fitz) para texto nativo; se for image-heavy, complementa
    com OCR/visão via Gemini multimodal (passa o PDF inteiro como Part).
  - Imagem: passa direto como Part para Gemini Vision.
  - DOCX/XLSX: extrai texto via python-docx/openpyxl.
  - Texto puro / CSV: lê direto.

Saída: `ExtractedAnexo` com:
  - `text_excerpt` (str): texto extraído resumido para o prompt
  - `gemini_parts` (list): Parts multimodais para passar ao Gemini Vertex
  - `mime` / `name`
"""
from __future__ import annotations

import asyncio
import base64
import io
import logging
import os
import re
from dataclasses import dataclass, field
from typing import Any, Optional
from urllib.parse import urlparse

import httpx

from xerticaproc.backend.models.copilot_schemas import Anexo

log = logging.getLogger(__name__)

MAX_TEXT_PER_DOC = 60_000
MAX_PDF_BYTES_FOR_GEMINI = 18 * 1024 * 1024  # 18 MB
PDF_TEXT_MIN_CHARS_PER_PAGE = 80  # abaixo disso considera image-heavy


@dataclass
class ExtractedAnexo:
    name: str
    mime: str
    text_excerpt: str = ""
    gemini_parts: list[Any] = field(default_factory=list)
    pages: int = 0
    bytes_size: int = 0
    notes: list[str] = field(default_factory=list)
    error: Optional[str] = None


# ─── helpers de I/O ──────────────────────────────────────────────────────────

def _guess_mime(name: str, default: str = "application/octet-stream") -> str:
    n = name.lower()
    if n.endswith(".pdf"):
        return "application/pdf"
    if n.endswith((".png",)):
        return "image/png"
    if n.endswith((".jpg", ".jpeg")):
        return "image/jpeg"
    if n.endswith(".webp"):
        return "image/webp"
    if n.endswith(".gif"):
        return "image/gif"
    if n.endswith(".docx"):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if n.endswith(".xlsx"):
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if n.endswith((".txt", ".md")):
        return "text/plain"
    if n.endswith(".csv"):
        return "text/csv"
    return default


async def _fetch_url(url: str) -> bytes:
    async with httpx.AsyncClient(timeout=60.0, follow_redirects=True) as cli:
        r = await cli.get(url)
        r.raise_for_status()
        return r.content


def _read_gcs_sync(gcs_uri: str) -> bytes:
    from google.cloud import storage  # type: ignore
    parsed = urlparse(gcs_uri)
    if parsed.scheme != "gs":
        raise ValueError(f"gcs_uri inválido: {gcs_uri}")
    bucket_name = parsed.netloc
    blob_name = parsed.path.lstrip("/")
    cli = storage.Client()
    return cli.bucket(bucket_name).blob(blob_name).download_as_bytes()


async def _read_anexo_bytes(anexo: Anexo) -> tuple[bytes, str]:
    """Retorna (bytes, mime). Suporta gcs_uri, url, ou inline texto."""
    name = anexo.nome or ""
    if anexo.gcs_uri:
        data = await asyncio.to_thread(_read_gcs_sync, anexo.gcs_uri)
        return data, _guess_mime(name)
    if anexo.url:
        data = await _fetch_url(anexo.url)
        # tenta deduzir mime pela URL
        return data, _guess_mime(anexo.url)
    if anexo.tipo == "texto":
        # Quando tipo=texto, o conteúdo vem inline em url (compat) ou nome.
        text = anexo.url or anexo.nome or ""
        return text.encode("utf-8"), "text/plain"
    raise ValueError(f"Anexo sem fonte (gcs_uri/url): {name}")


# ─── PDF via PyMuPDF ─────────────────────────────────────────────────────────

def _extract_pdf_pymupdf(data: bytes) -> tuple[str, int, bool]:
    """Retorna (texto, num_paginas, image_heavy_flag)."""
    try:
        import fitz  # PyMuPDF
    except Exception:
        log.warning("PyMuPDF não disponível; PDF cairá em multimodal Gemini")
        return "", 0, True
    try:
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception as e:
        log.warning("Falha abrindo PDF com PyMuPDF: %s", e)
        return "", 0, True
    parts: list[str] = []
    image_heavy_pages = 0
    n_pages = doc.page_count
    for i in range(n_pages):
        try:
            page = doc.load_page(i)
            txt = page.get_text("text") or ""
            txt = txt.strip()
            if len(txt) < PDF_TEXT_MIN_CHARS_PER_PAGE:
                image_heavy_pages += 1
            if txt:
                parts.append(f"--- página {i + 1} ---\n{txt}")
        except Exception:  # noqa: BLE001
            log.exception("Erro lendo página %d", i)
            image_heavy_pages += 1
    doc.close()
    full = "\n\n".join(parts)
    image_heavy = (n_pages > 0 and image_heavy_pages / max(n_pages, 1) >= 0.5)
    return full, n_pages, image_heavy


# ─── DOCX / XLSX ─────────────────────────────────────────────────────────────

def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document  # python-docx
    except Exception:
        return ""
    try:
        d = Document(io.BytesIO(data))
        out: list[str] = []
        for p in d.paragraphs:
            if p.text:
                out.append(p.text)
        for t in d.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                out.append(" | ".join(cells))
        return "\n".join(out)
    except Exception:  # noqa: BLE001
        log.exception("Falha extraindo DOCX")
        return ""


def _extract_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except Exception:
        return ""
    try:
        wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
        out: list[str] = []
        for ws in wb.worksheets:
            out.append(f"### planilha: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                vals = ["" if v is None else str(v) for v in row]
                if any(vals):
                    out.append(" | ".join(vals))
        return "\n".join(out)
    except Exception:  # noqa: BLE001
        log.exception("Falha extraindo XLSX")
        return ""


# ─── Gemini multimodal Parts ─────────────────────────────────────────────────

def _build_vertex_part(data: bytes, mime: str) -> Optional[Any]:
    """Cria um Part do Vertex SDK a partir de bytes."""
    try:
        from vertexai.generative_models import Part  # type: ignore
    except Exception:
        log.warning("vertexai indisponível; pulando Part multimodal")
        return None
    try:
        return Part.from_data(data=data, mime_type=mime)
    except Exception:  # noqa: BLE001
        log.exception("Falha criando Part Vertex")
        return None


# ─── Pipeline principal ──────────────────────────────────────────────────────

async def process_anexo(anexo: Anexo) -> ExtractedAnexo:
    out = ExtractedAnexo(name=anexo.nome or "anexo", mime="")
    try:
        data, mime = await _read_anexo_bytes(anexo)
    except Exception as e:  # noqa: BLE001
        log.warning("Falha lendo anexo %s: %s", anexo.nome, e)
        out.error = str(e)
        return out
    out.mime = mime
    out.bytes_size = len(data)

    if mime == "application/pdf":
        text, n_pages, image_heavy = await asyncio.to_thread(
            _extract_pdf_pymupdf, data,
        )
        out.pages = n_pages
        if text:
            out.text_excerpt = text[:MAX_TEXT_PER_DOC]
        # Sempre passa o PDF como Part para Gemini fazer multimodal
        # (cobre OCR de páginas escaneadas e leitura visual de tabelas/diagramas)
        if len(data) <= MAX_PDF_BYTES_FOR_GEMINI:
            part = _build_vertex_part(data, mime)
            if part is not None:
                out.gemini_parts.append(part)
                if image_heavy:
                    out.notes.append("PDF image-heavy: usando visão do Gemini")
        else:
            out.notes.append(f"PDF muito grande ({out.bytes_size} bytes); só texto")
        return out

    if mime.startswith("image/"):
        # Pillow apenas valida e normaliza
        try:
            from PIL import Image  # noqa: F401
            img = Image.open(io.BytesIO(data))
            img.verify()
        except Exception:  # noqa: BLE001
            out.notes.append("imagem inválida ou corrompida")
        part = _build_vertex_part(data, mime)
        if part is not None:
            out.gemini_parts.append(part)
        out.text_excerpt = f"[imagem anexada: {anexo.nome}]"
        return out

    if mime.endswith("wordprocessingml.document") or mime == "application/msword":
        text = await asyncio.to_thread(_extract_docx, data)
        out.text_excerpt = text[:MAX_TEXT_PER_DOC]
        return out

    if mime.endswith("spreadsheetml.sheet") or mime == "application/vnd.ms-excel":
        text = await asyncio.to_thread(_extract_xlsx, data)
        out.text_excerpt = text[:MAX_TEXT_PER_DOC]
        return out

    if mime.startswith("text/"):
        try:
            txt = data.decode("utf-8", errors="replace")
        except Exception:  # noqa: BLE001
            txt = ""
        out.text_excerpt = txt[:MAX_TEXT_PER_DOC]
        return out

    # fallback: tenta como texto
    try:
        txt = data.decode("utf-8", errors="replace")
        out.text_excerpt = txt[:MAX_TEXT_PER_DOC]
    except Exception:  # noqa: BLE001
        out.notes.append(f"mime {mime} não suportado para extração de texto")
    return out


async def process_anexos(
    anexos: Optional[list[Anexo]],
) -> list[ExtractedAnexo]:
    if not anexos:
        return []
    return await asyncio.gather(*(process_anexo(a) for a in anexos))


def render_anexos_for_prompt(extracted: list[ExtractedAnexo]) -> str:
    """Bloco markdown a injetar no prompt quando há anexos."""
    if not extracted:
        return ""
    lines: list[str] = ["\n## Anexos enviados pelo usuário\n"]
    for i, e in enumerate(extracted, start=1):
        header = (
            f"### Anexo {i}: {e.name}  "
            f"(mime={e.mime}, pages={e.pages}, bytes={e.bytes_size})"
        )
        lines.append(header)
        if e.error:
            lines.append(f"_erro: {e.error}_")
            continue
        if e.notes:
            lines.append("_notas: " + "; ".join(e.notes) + "_")
        if e.text_excerpt:
            lines.append("```\n" + e.text_excerpt[:MAX_TEXT_PER_DOC] + "\n```")
        elif e.gemini_parts:
            lines.append("_(conteúdo binário; analise a imagem/PDF anexado)_")
    lines.append(
        "\nUse essas informações para extrair fatos, decisões e atualizar o "
        "checklist. Em particular, identifique elementos de matriz de riscos, "
        "matriz de quantitativos, fontes de preço e requisitos técnicos."
    )
    return "\n".join(lines)


def collect_gemini_parts(extracted: list[ExtractedAnexo]) -> list[Any]:
    parts: list[Any] = []
    for e in extracted:
        parts.extend(e.gemini_parts)
    return parts
